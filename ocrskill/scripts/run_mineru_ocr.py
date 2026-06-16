#!/usr/bin/env python3
"""OCR PDFs, DOCX, and PPTX files with MinerU2.5-Pro.

The v2 output layout is one directory per source document:

  <output>/<docname>/
    <docname>.md
    <docname>_images.json
    <docname>-<hash>.jpg
    pages/page-0001.png
    json/page-0001.json
    layout/page-0001_layout.png
    manifest.json
    <docname>.docx
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import io
import json
import math
import re
import shutil
import subprocess
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable
from xml.sax.saxutils import escape

import fitz
import requests
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont

DEFAULT_MODEL_ID = "opendatalab/MinerU2.5-Pro-2605-1.2B"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".pptx"}
IMAGE_SCHEMA = "https://ezkotae.dev/schemas/images/v1"


@dataclass(frozen=True)
class DocumentJob:
    source_path: Path
    pdf_path: Path
    docname: str
    source_type: str
    output_dir: Path
    converted: bool = False


def to_jsonable(value: Any) -> Any:
    if isinstance(value, Path):
        return str(value)
    if isinstance(value, dict):
        return {str(k): to_jsonable(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [to_jsonable(v) for v in value]
    if hasattr(value, "item"):
        try:
            return value.item()
        except Exception:
            pass
    return value


def resolve_model_ref(model_id: str, local_files_only: bool) -> str:
    if Path(model_id).exists():
        return model_id
    if not local_files_only:
        return model_id
    try:
        from huggingface_hub import snapshot_download

        return snapshot_download(model_id, local_files_only=True)
    except Exception as exc:
        raise RuntimeError(
            f"Model is not available in the local Hugging Face cache: {model_id}. "
            "Download it first with snapshot_download or rerun with --allow-download."
        ) from exc


def collect_documents(input_path: Path, recursive: bool) -> list[Path]:
    if input_path.is_file():
        return [input_path] if input_path.suffix.lower() in SUPPORTED_EXTENSIONS else []

    pattern = "**/*" if recursive else "*"
    return sorted(
        path
        for path in input_path.glob(pattern)
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def find_libreoffice() -> str | None:
    for command in ("soffice", "libreoffice"):
        resolved = shutil.which(command)
        if resolved:
            return resolved

    candidates = [
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/usr/bin/soffice"),
        Path("/usr/bin/libreoffice"),
        Path("/usr/local/bin/soffice"),
        Path("/usr/local/bin/libreoffice"),
        Path("/opt/homebrew/bin/soffice"),
        Path("/opt/homebrew/bin/libreoffice"),
        Path("/usr/lib/libreoffice/program/soffice"),
        Path("/snap/bin/libreoffice"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def utc_now_z() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def converted_pdf_path(input_path: Path, output_dir: Path) -> Path:
    return output_dir / f"{input_path.stem}.pdf"


def newest_pdf(output_dir: Path) -> Path | None:
    candidates = sorted(output_dir.glob("*.pdf"), key=lambda path: path.stat().st_mtime, reverse=True)
    return candidates[0] if candidates else None


def convert_with_libreoffice(input_path: Path, output_dir: Path, libreoffice: str) -> Path:
    command = [
        libreoffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(input_path),
    ]
    result = subprocess.run(command, check=False, capture_output=True, text=True)
    if result.returncode != 0:
        details = (result.stderr or result.stdout).strip()
        raise RuntimeError(f"LibreOffice failed to convert {input_path}: {details}")

    converted = converted_pdf_path(input_path, output_dir)
    if converted.exists():
        return converted

    candidate = newest_pdf(output_dir)
    if candidate:
        return candidate
    raise RuntimeError(f"LibreOffice did not produce a PDF for {input_path}")


def convert_docx_with_word(input_path: Path, output_dir: Path) -> Path:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError("pywin32 is not installed; cannot use Microsoft Word COM fallback") from exc

    pdf_path = converted_pdf_path(input_path, output_dir)
    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(input_path),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        document.ExportAsFixedFormat(
            OutputFileName=str(pdf_path),
            ExportFormat=17,
            OpenAfterExport=False,
        )
    except Exception as exc:
        raise RuntimeError(f"Microsoft Word failed to convert {input_path}: {exc}") from exc
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    if pdf_path.exists():
        return pdf_path
    raise RuntimeError(f"Microsoft Word did not produce a PDF for {input_path}")


def convert_pptx_with_powerpoint(input_path: Path, output_dir: Path) -> Path:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise RuntimeError("pywin32 is not installed; cannot use Microsoft PowerPoint COM fallback") from exc

    pdf_path = converted_pdf_path(input_path, output_dir)
    pythoncom.CoInitialize()
    powerpoint = None
    presentation = None
    try:
        powerpoint = win32com.client.DispatchEx("PowerPoint.Application")
        presentation = powerpoint.Presentations.Open(
            str(input_path),
            ReadOnly=True,
            Untitled=False,
            WithWindow=False,
        )
        presentation.SaveAs(str(pdf_path), 32)
    except Exception as exc:
        raise RuntimeError(f"Microsoft PowerPoint failed to convert {input_path}: {exc}") from exc
    finally:
        if presentation is not None:
            try:
                presentation.Close()
            except Exception:
                pass
        if powerpoint is not None:
            try:
                powerpoint.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    if pdf_path.exists():
        return pdf_path
    raise RuntimeError(f"Microsoft PowerPoint did not produce a PDF for {input_path}")


def convert_with_office_com(input_path: Path, output_dir: Path) -> Path:
    suffix = input_path.suffix.lower()
    if sys.platform != "win32":
        raise RuntimeError("Microsoft Office COM fallback is only available on Windows")
    if suffix == ".docx":
        return convert_docx_with_word(input_path, output_dir)
    if suffix == ".pptx":
        return convert_pptx_with_powerpoint(input_path, output_dir)
    raise RuntimeError(f"Microsoft Office COM fallback does not support {suffix}")


def convert_to_pdf(input_path: Path, output_dir: Path) -> Path:
    suffix = input_path.suffix.lower()
    if suffix == ".pdf":
        return input_path
    if suffix not in {".docx", ".pptx"}:
        raise RuntimeError(f"Unsupported input type: {input_path}")

    output_dir.mkdir(parents=True, exist_ok=True)
    errors: list[str] = []
    libreoffice = find_libreoffice()
    if libreoffice:
        try:
            return convert_with_libreoffice(input_path, output_dir, libreoffice)
        except RuntimeError as exc:
            errors.append(str(exc))
    else:
        errors.append("LibreOffice/soffice was not found")

    try:
        return convert_with_office_com(input_path, output_dir)
    except RuntimeError as exc:
        errors.append(str(exc))

    raise RuntimeError(
        "Could not convert .docx/.pptx to PDF. Install LibreOffice, use Microsoft Office "
        f"with pywin32 on Windows, or provide a .pdf file. Details: {'; '.join(errors)}"
    )


def render_pdf_page(pdf_path: Path, page_index: int, image_path: Path, dpi: int) -> dict[str, Any]:
    with fitz.open(pdf_path) as pdf:
        page = pdf[page_index]
        matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        image_path.parent.mkdir(parents=True, exist_ok=True)
        pix.save(str(image_path))
        return {
            "page_width_pt": float(page.rect.width),
            "page_height_pt": float(page.rect.height),
            "image_width": int(pix.width),
            "image_height": int(pix.height),
        }


def render_all_pages(pdf_path: Path, output_dir: Path, dpi: int, workers: int) -> list[dict[str, Any]]:
    pages_dir = output_dir / "pages"
    with fitz.open(pdf_path) as pdf:
        page_count = len(pdf)

    def render_one(page_index: int) -> dict[str, Any]:
        page_number = page_index + 1
        page_name = f"page-{page_number:04d}"
        image_path = pages_dir / f"{page_name}.png"
        meta = render_pdf_page(pdf_path, page_index, image_path, dpi)
        return {
            "page_number": page_number,
            "page_name": page_name,
            "image_path": str(image_path),
            **meta,
        }

    if workers <= 1 or page_count <= 1:
        return [render_one(index) for index in range(page_count)]

    records: list[dict[str, Any]] = []
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = [executor.submit(render_one, index) for index in range(page_count)]
        for future in as_completed(futures):
            records.append(future.result())
    return sorted(records, key=lambda record: int(record["page_number"]))


class ApiRecognizer:
    def __init__(self, endpoint: str) -> None:
        endpoint = endpoint.rstrip("/")
        self.url = endpoint if endpoint.endswith("/ocr") else endpoint + "/ocr"

    def recognize(self, image_path: Path) -> dict[str, Any]:
        with image_path.open("rb") as handle:
            response = requests.post(
                self.url,
                files={"file": (image_path.name, handle, "image/png")},
                timeout=None,
            )
        if response.status_code >= 400:
            raise RuntimeError(
                f"OCR API failed for {image_path}: {response.status_code} {response.text[:1000]}"
            )
        return response.json()


class DirectRecognizer:
    def __init__(self, model_id: str, backend: str, image_analysis: bool, local_files_only: bool) -> None:
        try:
            from mineru_vl_utils import MinerUClient
            from mineru_vl_utils.post_process import json2md
        except ImportError as exc:
            raise RuntimeError(
                'Missing mineru-vl-utils. Install with: python -m pip install "mineru-vl-utils[transformers]"'
            ) from exc

        self.json2md = json2md
        model_ref = resolve_model_ref(model_id, local_files_only)
        if backend == "transformers":
            from transformers import AutoProcessor, Qwen2VLForConditionalGeneration

            model = Qwen2VLForConditionalGeneration.from_pretrained(
                model_ref,
                dtype="auto",
                device_map="auto",
                local_files_only=local_files_only,
            )
            processor = AutoProcessor.from_pretrained(model_ref, use_fast=True, local_files_only=local_files_only)
            self.client = MinerUClient(
                backend="transformers",
                model=model,
                processor=processor,
                image_analysis=image_analysis,
            )
        elif backend == "vllm-engine":
            from mineru_vl_utils import MinerULogitsProcessor
            from vllm import LLM

            llm = LLM(model=model_ref, logits_processors=[MinerULogitsProcessor])
            self.client = MinerUClient(
                backend="vllm-engine",
                vllm_llm=llm,
                image_analysis=image_analysis,
            )
        else:
            raise RuntimeError(f"Unsupported backend: {backend}")

    def recognize(self, image_path: Path) -> dict[str, Any]:
        with Image.open(image_path) as image:
            content_list = self.client.two_step_extract(image.convert("RGB"))
        return {
            "content_list": to_jsonable(content_list),
            "markdown": self.json2md(content_list),
        }


def text_from_item(item: dict[str, Any]) -> str:
    for key in ("text", "content", "markdown", "latex", "html"):
        value = item.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return ""


def content_items(node: Any) -> Iterable[dict[str, Any]]:
    if isinstance(node, list):
        for child in node:
            yield from content_items(child)
        return
    if not isinstance(node, dict):
        return

    has_text = bool(text_from_item(node))
    has_box = raw_bbox(node) is not None
    if has_text or has_box:
        yield node

    for key in ("children", "blocks", "lines", "spans", "cells", "items"):
        child = node.get(key)
        if isinstance(child, (dict, list)):
            yield from content_items(child)


def raw_bbox(item: dict[str, Any]) -> list[float] | None:
    for key in ("bbox", "box", "rect"):
        value = item.get(key)
        if isinstance(value, (list, tuple)) and len(value) >= 4:
            try:
                return [float(value[0]), float(value[1]), float(value[2]), float(value[3])]
            except (TypeError, ValueError):
                return None

    for key in ("poly", "polygon"):
        value = item.get(key)
        if not isinstance(value, (list, tuple)) or not value:
            continue
        points: list[tuple[float, float]] = []
        for point in value:
            if isinstance(point, (list, tuple)) and len(point) >= 2:
                try:
                    points.append((float(point[0]), float(point[1])))
                except (TypeError, ValueError):
                    pass
        if points:
            xs = [point[0] for point in points]
            ys = [point[1] for point in points]
            return [min(xs), min(ys), max(xs), max(ys)]
    return None


def normalize_bbox_to_pixels(
    bbox: list[float],
    image_width: int,
    image_height: int,
    page_width_pt: float,
    page_height_pt: float,
) -> tuple[float, float, float, float]:
    x0, y0, x1, y1 = bbox
    max_value = max(abs(value) for value in bbox)

    if max_value <= 1.5:
        x0, x1 = x0 * image_width, x1 * image_width
        y0, y1 = y0 * image_height, y1 * image_height
    elif max_value <= max(page_width_pt, page_height_pt) * 1.5:
        x_scale = image_width / page_width_pt if page_width_pt else 1.0
        y_scale = image_height / page_height_pt if page_height_pt else 1.0
        x0, x1 = x0 * x_scale, x1 * x_scale
        y0, y1 = y0 * y_scale, y1 * y_scale

    x0, x1 = sorted((max(0.0, x0), min(float(image_width), x1)))
    y0, y1 = sorted((max(0.0, y0), min(float(image_height), y1)))
    return x0, y0, x1, y1


def item_type(item: dict[str, Any]) -> str:
    value = item.get("type") or item.get("category") or item.get("label") or "block"
    return str(value)


def is_image_item(item: dict[str, Any]) -> bool:
    return item_type(item).lower() == "image"


def draw_layout_image(page_record: dict[str, Any], layout_path: Path) -> int:
    content_list = page_record.get("content_list", [])
    image_path = Path(page_record["image_path"])
    with Image.open(image_path) as source:
        image = source.convert("RGB")

    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    colors = {
        "title": (220, 38, 38),
        "text": (37, 99, 235),
        "table": (5, 150, 105),
        "formula": (147, 51, 234),
        "image": (234, 88, 12),
        "block": (75, 85, 99),
    }
    count = 0

    for item in content_items(content_list):
        bbox = raw_bbox(item)
        if not bbox:
            continue
        box = normalize_bbox_to_pixels(
            bbox,
            int(page_record["image_width"]),
            int(page_record["image_height"]),
            float(page_record["page_width_pt"]),
            float(page_record["page_height_pt"]),
        )
        if box[2] - box[0] < 2 or box[3] - box[1] < 2:
            continue

        kind = item_type(item).lower()
        color = next((value for name, value in colors.items() if name in kind), colors["block"])
        draw.rectangle(box, outline=color, width=3)
        label = item_type(item)[:28]
        text_box = draw.textbbox((box[0] + 2, box[1] + 2), label, font=font)
        draw.rectangle(text_box, fill=(255, 255, 255))
        draw.text((box[0] + 2, box[1] + 2), label, fill=color, font=font)
        count += 1

    layout_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(layout_path)
    return count


def fallback_markdown(content_list: Any) -> str:
    lines: list[str] = []
    for item in content_items(content_list):
        text = text_from_item(item)
        if text:
            lines.append(text)
    return "\n\n".join(lines)


def configure_section(section: Any, width_pt: float, height_pt: float) -> None:
    section.page_width = Pt(width_pt)
    section.page_height = Pt(height_pt)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    section.header_distance = Pt(0)
    section.footer_distance = Pt(0)


def textbox_xml(shape_id: str, x_pt: float, y_pt: float, width_pt: float, height_pt: float, text: str) -> str:
    font_size = max(6, min(12, height_pt / max(1, math.ceil(len(text) / 50)) / 1.6))
    half_points = int(round(font_size * 2))
    escaped_lines = [escape(line) for line in text.splitlines() or [""]]
    text_runs = []
    for index, line in enumerate(escaped_lines):
        if index:
            text_runs.append("<w:br/>")
        text_runs.append(f'<w:t xml:space="preserve">{line}</w:t>')
    text_body = "".join(text_runs)
    return f"""
<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:v="urn:schemas-microsoft-com:vml"
        xmlns:o="urn:schemas-microsoft-com:office:office">
  <v:shape id="{shape_id}" type="#_x0000_t202"
    style="position:absolute;margin-left:{x_pt:.2f}pt;margin-top:{y_pt:.2f}pt;width:{width_pt:.2f}pt;height:{height_pt:.2f}pt;z-index:1;mso-position-horizontal-relative:page;mso-position-vertical-relative:page"
    stroked="f" filled="f">
    <v:textbox inset="0,0,0,0">
      <w:txbxContent>
        <w:p>
          <w:r>
            <w:rPr><w:sz w:val="{half_points}"/></w:rPr>
            {text_body}
          </w:r>
        </w:p>
      </w:txbxContent>
    </v:textbox>
  </v:shape>
</w:pict>
"""


def add_textbox(doc: Document, shape_id: str, x_pt: float, y_pt: float, width_pt: float, height_pt: float, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run()
    run._r.append(parse_xml(textbox_xml(shape_id, x_pt, y_pt, width_pt, height_pt, text)))


VML_IMAGE_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
    "o": "urn:schemas-microsoft-com:office:office",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def _ns_attrs(ns_map: dict[str, str]) -> str:
    return " ".join(f'xmlns:{key}="{value}"' for key, value in ns_map.items())


def _image_vml(shape_id: str, relationship_id: str, x_pt: float, y_pt: float, w_pt: float, h_pt: float) -> str:
    return (
        f'<w:pict {_ns_attrs(VML_IMAGE_NS)}>'
        f'<v:shape id="{shape_id}" type="#_x0000_t75" '
        f'style="position:absolute;margin-left:{x_pt:.2f}pt;margin-top:{y_pt:.2f}pt;'
        f"width:{w_pt:.2f}pt;height:{h_pt:.2f}pt;z-index:2;"
        f'mso-position-horizontal-relative:page;mso-position-vertical-relative:page" '
        f'stroked="f" filled="f">'
        f'<v:imagedata r:id="{relationship_id}" o:title="img"/>'
        f"</v:shape></w:pict>"
    )


class ImageExtractor:
    def __init__(self, docname: str, output_dir: Path, image_quality: int) -> None:
        self.docname = docname
        self.output_dir = output_dir
        self.image_quality = image_quality
        self.images: dict[str, str] = {}

    def extract(self, page_records: list[dict[str, Any]]) -> dict[str, str]:
        for record in page_records:
            extracted: list[dict[str, Any]] = []
            image_path = Path(record["image_path"])
            if not image_path.exists():
                record["extracted_images"] = extracted
                continue

            with Image.open(image_path) as page_image:
                page = page_image.convert("RGB")
                for item in content_items(record.get("content_list", [])):
                    if not is_image_item(item):
                        continue
                    bbox = raw_bbox(item)
                    if not bbox:
                        continue
                    x0, y0, x1, y1 = normalize_bbox_to_pixels(
                        bbox,
                        int(record["image_width"]),
                        int(record["image_height"]),
                        float(record["page_width_pt"]),
                        float(record["page_height_pt"]),
                    )
                    if x1 - x0 < 4 or y1 - y0 < 4:
                        continue

                    cropped = page.crop((int(x0), int(y0), int(x1), int(y1))).convert("RGB")
                    buffer = io.BytesIO()
                    cropped.save(buffer, format="JPEG", quality=self.image_quality, optimize=True)
                    image_bytes = buffer.getvalue()
                    hash10 = hashlib.sha256(image_bytes).hexdigest()[:10]
                    key = f"{self.docname}-{hash10}.jpg"
                    output_path = self.output_dir / key
                    output_path.write_bytes(image_bytes)
                    self.images[key] = base64.b64encode(image_bytes).decode("ascii")
                    extracted.append(
                        {
                            "key": key,
                            "path": str(output_path),
                            "bbox_pixels": [x0, y0, x1, y1],
                        }
                    )
            record["extracted_images"] = extracted
        return self.images


IMAGE_REF_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def rewrite_md_image_refs(md_text: str, image_keys: list[str]) -> str:
    used_keys: list[str] = []
    key_iter = iter(image_keys)

    def replace(match: re.Match[str]) -> str:
        try:
            key = next(key_iter)
        except StopIteration:
            return match.group(0)
        used_keys.append(key)
        return f"![{match.group(1)}]({key})"

    rewritten = IMAGE_REF_RE.sub(replace, md_text)
    missing = [key for key in image_keys if key not in used_keys]
    if missing:
        suffix = "\n\n" + "\n\n".join(f"![]({key})" for key in missing)
        rewritten = rewritten.rstrip() + suffix
    return rewritten


def build_images_json(images: dict[str, str], docname: str, source_type: str) -> dict[str, Any]:
    return {
        "$schema": IMAGE_SCHEMA,
        "docname": docname,
        "images": images,
        "generated_by": "ocrskill-v2",
        "generated_at": utc_now_z(),
        "source_type": source_type,
        "image_count": len(images),
    }


def recognize_page_batch(
    page_metas: list[dict[str, Any]],
    output_dir: Path,
    recognizer: ApiRecognizer | DirectRecognizer,
    workers: int,
    skip_existing: bool,
    job: DocumentJob,
) -> list[dict[str, Any]]:
    def recognize_one(meta: dict[str, Any]) -> dict[str, Any]:
        page_name = str(meta["page_name"])
        image_path = Path(meta["image_path"])
        json_path = output_dir / "json" / f"{page_name}.json"
        layout_path = output_dir / "layout" / f"{page_name}_layout.png"

        if skip_existing and json_path.exists():
            with json_path.open("r", encoding="utf-8") as handle:
                record = json.load(handle)
            record.update(
                {
                    "source_file": str(job.source_path),
                    "source_name": job.source_path.name,
                    "source_type": job.source_type,
                    "source_pdf": str(job.pdf_path),
                    "source_pdf_name": job.pdf_path.name,
                    "docname": job.docname,
                    "page_number": meta["page_number"],
                    "page_name": page_name,
                    "image_path": str(image_path),
                    "layout_path": str(layout_path),
                    "converted": job.converted,
                }
            )
        else:
            ocr_result = recognizer.recognize(image_path)
            record = {
                "source_file": str(job.source_path),
                "source_name": job.source_path.name,
                "source_type": job.source_type,
                "source_pdf": str(job.pdf_path),
                "source_pdf_name": job.pdf_path.name,
                "docname": job.docname,
                "page_number": meta["page_number"],
                "page_name": page_name,
                "image_path": str(image_path),
                "layout_path": str(layout_path),
                "converted": job.converted,
                "content_list": ocr_result.get("content_list", ocr_result),
                "markdown": ocr_result.get("markdown", ""),
                "page_width_pt": meta["page_width_pt"],
                "page_height_pt": meta["page_height_pt"],
                "image_width": meta["image_width"],
                "image_height": meta["image_height"],
            }
            json_path.parent.mkdir(parents=True, exist_ok=True)
            with json_path.open("w", encoding="utf-8") as handle:
                json.dump(to_jsonable(record), handle, ensure_ascii=False, indent=2)

        record["layout_box_count"] = draw_layout_image(record, layout_path)
        return record

    if isinstance(recognizer, ApiRecognizer) and workers > 1 and len(page_metas) > 1:
        records: list[dict[str, Any]] = []
        with ThreadPoolExecutor(max_workers=workers) as executor:
            futures = [executor.submit(recognize_one, meta) for meta in page_metas]
            for future in as_completed(futures):
                records.append(future.result())
        return sorted(records, key=lambda record: int(record["page_number"]))

    return [recognize_one(meta) for meta in page_metas]


def write_markdown_file(page_records: list[dict[str, Any]], output_dir: Path, docname: str) -> Path:
    parts = [f"# {docname}", ""]
    for record in sorted(page_records, key=lambda item: int(item["page_number"])):
        markdown = record.get("markdown") or fallback_markdown(record.get("content_list", []))
        image_keys = [item["key"] for item in record.get("extracted_images", [])]
        markdown = rewrite_md_image_refs(markdown, image_keys)
        record["markdown_rewritten"] = markdown
        parts.extend([f"## Page {record['page_number']}", "", markdown, ""])

    md_path = output_dir / f"{docname}.md"
    md_path.write_text("\n".join(parts).strip() + "\n", encoding="utf-8")
    return md_path


def write_docx(page_records: list[dict[str, Any]], output_path: Path, mode: str) -> None:
    doc = Document()
    if not page_records:
        doc.add_paragraph("No OCR pages were generated.")
        doc.save(output_path)
        return

    shape_index = 1000
    first_page = True
    total_images = 0

    for record in sorted(page_records, key=lambda item: int(item["page_number"])):
        if first_page:
            section = doc.sections[0]
            first_page = False
        else:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
        width_pt = float(record["page_width_pt"])
        height_pt = float(record["page_height_pt"])
        img_w = int(record["image_width"])
        img_h = int(record["image_height"])
        configure_section(section, width_pt, height_pt)

        markdown = record.get("markdown_rewritten") or record.get("markdown") or fallback_markdown(record.get("content_list", []))

        if mode == "image-plus-text":
            doc.add_picture(record["image_path"], width=Pt(width_pt))
            doc.add_page_break()
            doc.add_paragraph(f'{record["source_name"]} - page {record["page_number"]}')
            for paragraph_text in markdown.split("\n\n"):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            continue

        if mode == "markdown":
            doc.add_paragraph(f'{record["source_name"]} - page {record["page_number"]}')
            for paragraph_text in markdown.split("\n\n"):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
            continue

        image_entries = iter(record.get("extracted_images", []))
        added = 0

        for item in content_items(record.get("content_list", [])):
            kind = item_type(item).lower()
            text = text_from_item(item)
            bbox = raw_bbox(item)
            if not bbox:
                continue

            if mode == "mixed" and kind == "image":
                try:
                    image_entry = next(image_entries)
                except StopIteration:
                    continue
                x0, y0, x1, y1 = image_entry["bbox_pixels"]
                x_pt = x0 / img_w * width_pt
                y_pt = y0 / img_h * height_pt
                w_pt = (x1 - x0) / img_w * width_pt
                h_pt = (y1 - y0) / img_h * height_pt

                try:
                    with Path(image_entry["path"]).open("rb") as image_file:
                        r_id, _ = doc.part.get_or_add_image(image_file)
                except Exception:
                    continue

                vml = _image_vml(f"_mineru_img_{shape_index}", r_id, x_pt, y_pt, w_pt, h_pt)
                shape_index += 1
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run()
                try:
                    run._r.append(parse_xml(vml))
                    total_images += 1
                    added += 1
                except Exception:
                    pass
                continue

            if not text:
                continue

            x0, y0, x1, y1 = normalize_bbox_to_pixels(bbox, img_w, img_h, width_pt, height_pt)
            x_pt_t = x0 / img_w * width_pt
            y_pt_t = y0 / img_h * height_pt
            w_pt_t = max(4.0, (x1 - x0) / img_w * width_pt)
            h_pt_t = max(4.0, (y1 - y0) / img_h * height_pt)
            add_textbox(doc, f"_mineru_txt_{shape_index}", x_pt_t, y_pt_t, w_pt_t, h_pt_t, text)
            shape_index += 1
            added += 1

        if added == 0:
            doc.add_paragraph(markdown or f'{record["source_name"]} - page {record["page_number"]}')

    if mode == "mixed" and total_images:
        print(f"  Embedded {total_images} cropped images in DOCX")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def process_document(
    job: DocumentJob,
    recognizer: ApiRecognizer | DirectRecognizer,
    dpi: int,
    workers: int,
    skip_existing: bool,
    no_docx: bool,
    docx_mode: str,
    image_quality: int,
) -> dict[str, Any]:
    job.output_dir.mkdir(parents=True, exist_ok=True)
    page_metas = render_all_pages(job.pdf_path, job.output_dir, dpi, workers)
    page_records = recognize_page_batch(
        page_metas,
        job.output_dir,
        recognizer,
        workers,
        skip_existing,
        job,
    )

    extractor = ImageExtractor(job.docname, job.output_dir, image_quality)
    images = extractor.extract(page_records)

    md_path = write_markdown_file(page_records, job.output_dir, job.docname)
    images_json_path = job.output_dir / f"{job.docname}_images.json"
    images_json_path.write_text(
        json.dumps(build_images_json(images, job.docname, job.source_type), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    docx_path: Path | None = None
    if not no_docx:
        docx_path = job.output_dir / f"{job.docname}.docx"
        write_docx(page_records, docx_path, docx_mode)

    manifest = {
        "created_at": utc_now_z(),
        "generated_by": "ocrskill-v2",
        "source_file": str(job.source_path),
        "source_type": job.source_type,
        "source_pdf": str(job.pdf_path),
        "converted": job.converted,
        "docname": job.docname,
        "model_id": DEFAULT_MODEL_ID,
        "dpi": dpi,
        "workers": workers,
        "image_format": "jpg",
        "image_quality": image_quality,
        "docx_mode": None if no_docx else docx_mode,
        "page_count": len(page_records),
        "image_count": len(images),
        "markdown_path": str(md_path),
        "images_json_path": str(images_json_path),
        "docx_path": str(docx_path) if docx_path else None,
        "pages_dir": str(job.output_dir / "pages"),
        "json_dir": str(job.output_dir / "json"),
        "layout_dir": str(job.output_dir / "layout"),
    }
    (job.output_dir / "manifest.json").write_text(
        json.dumps(to_jsonable(manifest), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return manifest


def build_jobs(documents: list[Path], output_root: Path, temp_root: Path) -> list[DocumentJob]:
    jobs: list[DocumentJob] = []
    convert_root = temp_root / "converted"
    for source_path in documents:
        source_type = source_path.suffix.lower().lstrip(".")
        docname = source_path.stem
        output_dir = output_root / docname
        pdf_path = convert_to_pdf(source_path, convert_root) if source_type != "pdf" else source_path
        jobs.append(
            DocumentJob(
                source_path=source_path,
                pdf_path=pdf_path,
                docname=docname,
                source_type=source_type,
                output_dir=output_dir,
                converted=source_type != "pdf",
            )
        )
    return jobs


def make_recognizer(args: argparse.Namespace) -> tuple[ApiRecognizer | DirectRecognizer, str]:
    if args.endpoint or args.backend == "api":
        if not args.endpoint:
            raise RuntimeError("--endpoint is required when --backend api is used.")
        return ApiRecognizer(args.endpoint), f"api:{args.endpoint}"

    return (
        DirectRecognizer(args.model_id, args.backend, args.image_analysis, not args.allow_download),
        args.backend,
    )


def main() -> None:
    parser = argparse.ArgumentParser(description="OCR PDFs, DOCX, and PPTX files with MinerU2.5-Pro.")
    parser.add_argument("input", nargs="?", default=".", help="PDF/DOCX/PPTX file or folder.")
    parser.add_argument("--output", help="Output root. Defaults to <file-parent> or <folder>/mineru_ocr_output.")
    parser.add_argument("--model-id", default=DEFAULT_MODEL_ID)
    parser.add_argument("--backend", default="transformers", choices=["transformers", "vllm-engine", "api"])
    parser.add_argument("--endpoint", help="Existing OCR endpoint, e.g. the endpoint printed by launch_mineru_api.py.")
    parser.add_argument("--dpi", type=int, default=220)
    parser.add_argument("--workers", type=int, default=1)
    parser.add_argument("--recursive", action="store_true")
    parser.add_argument("--skip-existing", action="store_true")
    parser.add_argument("--no-docx", action="store_true")
    parser.add_argument("--docx-mode", default="mixed", choices=["mixed", "textboxes", "image-plus-text", "markdown"])
    parser.add_argument("--image-quality", type=int, default=92, help="JPEG quality for extracted image crops.")
    parser.add_argument("--image-analysis", action="store_true")
    parser.add_argument("--allow-download", action="store_true", help="Allow Hugging Face network checks/downloads at load time.")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    documents = collect_documents(input_path, args.recursive)
    if not documents:
        print(f"No supported documents found at {input_path} (.pdf, .docx, .pptx).")
        return

    if args.workers < 1:
        raise RuntimeError("--workers must be at least 1.")
    if not 1 <= args.image_quality <= 100:
        raise RuntimeError("--image-quality must be between 1 and 100.")

    if args.output:
        output_root = Path(args.output).resolve()
    elif input_path.is_file():
        output_root = input_path.parent
    else:
        output_root = input_path / "mineru_ocr_output"
    output_root.mkdir(parents=True, exist_ok=True)

    run_manifests: list[dict[str, Any]] = []

    with tempfile.TemporaryDirectory(prefix="ocrskill-v2-") as temp_dir:
        jobs = build_jobs(documents, output_root, Path(temp_dir))
        recognizer, recognizer_mode = make_recognizer(args)
        for job in jobs:
            print(f"Processing {job.source_path.name} -> {job.output_dir}")
            manifest = process_document(
                job,
                recognizer,
                args.dpi,
                args.workers,
                args.skip_existing,
                args.no_docx,
                args.docx_mode,
                args.image_quality,
            )
            manifest["recognizer"] = recognizer_mode
            manifest["model_id"] = args.model_id
            (job.output_dir / "manifest.json").write_text(
                json.dumps(to_jsonable(manifest), ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            run_manifests.append(manifest)

    print(f"Processed {len(run_manifests)} document(s) into {output_root}")


if __name__ == "__main__":
    try:
        main()
    except RuntimeError as exc:
        print(f"Error: {exc}")
        raise SystemExit(1) from exc
